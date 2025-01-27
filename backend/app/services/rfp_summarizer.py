# from transformers import pipeline

# summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# async def process_rfp(text):
#     # You may want to implement a more sophisticated summarization method
#     summary = summarizer(text, max_length=150, min_length=50, do_sample=False)[0]['summary_text']
#     return summary



import os
from typing import Optional, List
import uuid
from PyPDF2 import PdfReader
import openai
from docx import Document


def add_table(header_row, data_rows, document):
    table = document.add_table(rows=1, cols=len(header_row))
    table.style = 'Table Grid'
    # Add header row
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(header_row):
        hdr_cells[idx].text = text.strip()
    # Add data rows
    for row in data_rows:
        row_cells = table.add_row().cells
        for idx, text in enumerate(row):
            row_cells[idx].text = text.strip()

# Function to convert a .md file to .docx
def convert_md_to_docx(md_file_path, docx_output_path):
    # Read the Markdown file
    with open(md_file_path, "r", encoding="utf-8") as file:
        md_content = file.read()

    # Create a Word document
    doc = Document()

    # Parse the Markdown content
    lines = md_content.split("\n")

    # Temporary storage for table data
    header = []
    rows = []

    # Process lines from the Markdown content
    for line in lines:
        if line.startswith("**") and line.endswith("**"):
            if header:  # Finalize any pending table
                add_table(header, rows, doc)
                header, rows = [], []
            doc.add_heading(line.strip("**"), level=2)
        elif line.startswith("* "):
            if header:  # Finalize any pending table
                add_table(header, rows, doc)
                header, rows = [], []
            doc.add_paragraph(line.strip("* "), style='List Bullet')
        elif line.startswith("|"):
            cols = line.strip("|").split("|")
            if "---" in line:  # Ignore table header separators
                continue
            elif not header:  # Initialize header row
                header = cols
            else:  # Add data rows
                rows.append(cols)
        elif line.strip() == "":
            if header:  # Finalize any pending table
                add_table(header, rows, doc)
                header, rows = [], []
            continue
        else:
            if header:  # Finalize any pending table
                add_table(header, rows, doc)
                header, rows = [], []
            doc.add_paragraph(line)

    # Finalize any remaining table
    if header:
        add_table(header, rows, doc)

    # Save the document
    doc.save(docx_output_path)
    print(f"Document saved to: {docx_output_path}")

# Local model API
def local_model_api(inputs: str) -> str:
    api_url = "http://192.168.3.202:1234/v1"  
    api_key = "lm_studio"  
    openai.api_base = api_url  
    openai.api_key = api_key
    try:
        response = openai.ChatCompletion.create(
            model="llama-3.2-3b-instruct",  
            messages=[{"role": "user", "content": inputs}],
            temperature=0.2,
            top_p=1
        )
        return response['choices'][0]['message']['content'].strip() 
    except Exception as e:
        return f"Error querying the LLM: {e}"

class LMStudioLLM:
    def _call(self, prompt: str, stop: Optional[List[str]] = None) -> str:
        return local_model_api(prompt)

    @property
    def _identifying_params(self) -> dict:
        return {}

    @property
    def _llm_type(self) -> str:
        return "lm_studio"

def extract_text_from_pdf(pdf_path: str) -> str:
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            text += page_text or ""  
        return text
    except Exception as e:
        return f"Error extracting text from PDF: {e}"

def generate_rfp_prompt(text: str) -> str:
    prompt = f"""
Please extract the following things from the PDF text:
- Current Challenges (Current issues or work or summary of existing solutions)
- Scope of Work/Statement of Work/Statement of Services
- Evaluation Criteria (generate answer in table format with criteria and points)
- Forms/Documentation Required
- Compliance (certificate, qualification)
- Project Goals
- Additional Features (Future Scalability)
- Important Dates or schedule including Question and answer date and must include submission date or due date (consider all the dates, in a table format)
- Submission process or method (Proposal contains, response method ,submission requirements)
- Please include Mandatory or must-have requirements mentioned in the document.
Here is the PDF text:
{text}
Please include Mandatory or must-have requirements mentioned in the document.
"""
    return prompt

def save_summary_to_files(summary: str, output_dir: str):
    # Generate a unique filename for the .docx file
    docx_filename = f"RFP_Summary_{uuid.uuid4().hex}.docx"
    markdown_path = os.path.join(output_dir, f"RFP_Summary_{uuid.uuid4().hex}.md")  # Create unique markdown filename
    docx_output_path = os.path.join(output_dir, docx_filename)

    try:
        # Save the markdown file
        with open(markdown_path, "w") as md_file:
            md_file.write(summary)

        # Convert Markdown to DOCX
        convert_md_to_docx(markdown_path, docx_output_path)

        return docx_filename  # Return the dynamically generated filename
    except Exception as e:
        print(f"Error saving summary: {e}")
        return None

def process_rfp(pdf_path: str, output_dir: str):
    try:
        rfp_text = extract_text_from_pdf(pdf_path)
        if not rfp_text.strip():
            raise ValueError("No text extracted from the PDF.")
        prompt = generate_rfp_prompt(rfp_text)

        llm = LMStudioLLM()
        summary = llm._call(prompt)

        # Save the summary to files and return the generated .docx filename
        docx_filename = save_summary_to_files(summary, output_dir)
        
        return summary, docx_filename
    except Exception as e:
        print(f"Error processing RFP: {e}")
        return None, None

if __name__ == "__main__":
    pdf_path = "approches/RFP Documents/1257048-event.pdf"  # Adjust path as needed
    output_dir = "output"

    os.makedirs(output_dir, exist_ok=True)

    process_rfp(pdf_path, output_dir)