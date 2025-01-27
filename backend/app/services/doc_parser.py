import re
import os
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFSyntaxError
from pdfminer.psexceptions import PSSyntaxError
from docx import Document

class DocumentParser:
    def __init__(self, filepath):
        self.filepath = filepath

    def get_cleaned_text(self):
        ext = os.path.splitext(self.filepath)[-1].lower()
        text = ""

        try:
            if ext == '.pdf':
                text = self.extract_pdf_text()
            elif ext == '.docx':
                text = self.convert_docx_to_text()
            elif ext == '.doc':
                text = self.convert_doc_to_text()
            elif ext == '.txt':
                with open(self.filepath, 'r') as file:
                    text = file.read()

            text = self.clean_text(text)
        except Exception as e:
            print(f"\nError processing {self.filepath}: {e}")
        
        return text

    def extract_pdf_text(self):
        try:
            text = extract_text(self.filepath)
            return text
        except (PDFSyntaxError, PSSyntaxError) as e:
            print(f"\nError extracting text from {self.filepath}: {e}")
            return ""

    def convert_docx_to_text(self):
        try:
            doc = Document(self.filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"\nError extracting text from {self.filepath}: {e}")
            return ""

    def convert_doc_to_text(self):
        try:
            docx_text = self.convert_doc_to_docx(self.filepath)
            return docx_text
        except Exception as e:
            print(f"\nError converting .doc to .docx for {self.filepath}: {e}")
            return ""

    @staticmethod
    def convert_doc_to_docx(doc_file):
        doc = Document(doc_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text

    @staticmethod
    def clean_text(text):
        text = re.sub(r"[^\w\s/#@._+\-/\\-]", ' ', text)  # Remove special characters
        text = re.sub(r"\s{2,}", " ", text)  # Replace multiple spaces with a single space
        text = re.sub(r"[•\t▪➢❖]", '', text)  # Remove specific bullet points and characters
        text = text.strip()
        return text.lower()

# if __name__ == '__main__':
    # main('/home/alois/Abhiraj/6. Resume Parser Model/og_respa_data_collection/test_resumes')


