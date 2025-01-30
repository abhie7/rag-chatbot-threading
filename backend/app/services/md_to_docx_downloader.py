import os
from docx import Document
from fastapi import HTTPException
import aiofiles

async def add_table(header_row, data_rows, document):
    table = document.add_table(rows=1, cols=len(header_row))
    table.style = 'Table Grid'
    # Add header row
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(header_row):
        hdr_cells[idx].text = text.strip()
    # Add data rows
    for row in data_rows:
        row_cells = table.add_row().cells
        for idx, text in enumerate(row):
            row_cells[idx].text = text.strip()

# Function to convert a .md file to .docx
async def convert_md_to_docx(md_file_path, docx_output_path):
    try:
        # Read the Markdown file
        async with aiofiles.open(md_file_path, "r", encoding="utf-8") as file:
            md_content = await file.read()

        # Create a Word document
        doc = Document()

        # Parse the Markdown content
        lines = md_content.split("\n")

        # Temporary storage for table data
        header = []
        rows = []

        # Process lines from the Markdown content
        for line in lines:
            if line.startswith("**") and line.endswith("**"):
                if header:  # Finalize any pending table
                    await add_table(header, rows, doc)
                    header, rows = [], []
                doc.add_heading(line.strip("**"), level=2)
            elif line.startswith("* "):
                if header:  # Finalize any pending table
                    await add_table(header, rows, doc)
                    header, rows = [], []
                doc.add_paragraph(line.strip("* "), style='List Bullet')
            elif line.startswith("|"):
                cols = line.strip("|").split("|")
                if "---" in line:  # Ignore table header separators
                    continue
                elif not header:  # Initialize header row
                    header = cols
                else:  # Add data rows
                    rows.append(cols)
            elif line.strip() == "":
                if header:  # Finalize any pending table
                    await add_table(header, rows, doc)
                    header, rows = [], []
                continue
            else:
                if header:  # Finalize any pending table
                    await add_table(header, rows, doc)
                    header, rows = [], []
                doc.add_paragraph(line)

        # Finalize any remaining table
        if header:
            await add_table(header, rows, doc)

        # Save the document
        doc.save(docx_output_path)
        print(f"Document saved to: {docx_output_path}")
    except Exception as e:
        print(f"Error converting Markdown to DOCX: {e}")
        raise HTTPException(status_code=500, detail="Error converting Markdown to DOCX")

async def save_summary_to_files(summary: str, output_dir: str, filename: str = None, document_hash: str = None):
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Generate a unique filename for the .docx file
    docx_filename = f"Summary_{filename}_{document_hash}.docx"
    markdown_path = os.path.join(output_dir, f"Summary_{filename}_{document_hash}.md")
    docx_output_path = os.path.join(output_dir, docx_filename)

    try:
        # Save the markdown file
        async with aiofiles.open(markdown_path, "w", encoding="utf-8") as md_file:
            await md_file.write(summary)

        # Convert Markdown to DOCX
        await convert_md_to_docx(markdown_path, docx_output_path)

        return docx_output_path  # Return the path to the generated .docx file
    except Exception as e:
        print(f"Error saving summary: {e}")
        raise HTTPException(status_code=500, detail="Error saving summary")